import pandas as pd
from datetime import datetime


data = {
    'Date': ['2023-02-01', '2023-03-01', '2023-04-01'],
    'Sales': [1000, 1200, 1500],
    'Expenses': [500, 600, 700],
    'Profit': [500, 600, 800]
}


df = pd.DataFrame(data)


df['Date'] = pd.to_datetime(df['Date'])


df['Monthly Profit'] = df.groupby('Date').transform(lambda x: x[-1] - x.iloc[:-2].sum())


print("Accounting Report:")
print(df)


from docx import Document

doc = Document()

for index, row in df.iterrows():
    paragraph = doc.add_paragraph()
    run = paragraph.add_run(f"Date: {row['Date']}\n")
    run.font.size = pt=12
    
    run = paragraph.add_run(f"Sales: {row['Sales']}\n")
    run.font.size = pt=12
    
    run = paragraph.add_run(f"Expenses: {row['Expenses']}\n")
    run.font.size = pt=12
    
    run = paragraph.add_run(f"Profit: {row['Monthly Profit']}\n")
    run.font.size = pt=12

doc.save("accounting_report.docx")